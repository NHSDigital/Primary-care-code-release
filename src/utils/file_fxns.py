import os
import numpy as np
import pandas as pd
import zipfile
import docx
from pathlib import PurePath
from email.mime.multipart import MIMEMultipart
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from email.mime.base import MIMEBase
from email import encoders, generator
import logging
import openpyxl
from openpyxl.utils.cell import get_column_letter 

#---------------------------------------------------------------------------------------------

def txt_save(df:pd.DataFrame, filename:str, filepath:str):
    """
    Function Actions:
    - Saves a pandas dataframe to a txt file using tab spacing.
    - Prints a message with filename and file location to the terminal.
    """
    df.to_csv(PurePath(filepath, filename), index = False, sep ='\t')
    logging.info(f'{filename} saved to {str(filepath)}')

#---------------------------------------------------------------------------------------------

def csv_save(df:pd.DataFrame, filename:str, filepath:str):
    """
    Function Actions:
    - Saves a pandas dataframe to a csv file using ansi encoding.
    - Prints a message with filename and file location to the terminal.
    """
    df.to_csv(PurePath(filepath, filename), index = False, encoding='ansi')
    logging.info(f'{filename} saved to {str(filepath)}')

#---------------------------------------------------------------------------------------------

def excel_save(df, filename:str, filepath:str):
    """
    Function Actions:
    - Saves a pandas dataframe to an excel file.
    - Prints a message with filename and file location to the terminal.
    """
    df.to_excel(PurePath(filepath, filename), index = False)
    logging.info(f'{filename} saved to {str(filepath)}')

#---------------------------------------------------------------------------------------------
 
def zip_files(content_filepath_name:list, zf_filepath_name:dict):
    """
    Function Actions:
    - Loops through content_filepath_name: a list of dictionaries containing the contents of the zipfile - a file location and a file name.
    - Adds each file to the zip file.
    - Saves the zip file using the name and filepath specified in the zf_filepath_name dictionary.
    - Prints a message with zip filename and file location to the terminal.
    """

    files = []

    with zipfile.ZipFile(PurePath(zf_filepath_name['filepath'], zf_filepath_name['name']), 'w',
                     compression=zipfile.ZIP_DEFLATED,
                     compresslevel=6) as zf:
        
        # file locations and names inside zip file
        for i in content_filepath_name:
            zf.write(PurePath(i['filepath'], i['name']), arcname=i['name'])
            files.append(i['name'])
    
    files = " and ".join(files)
    logging.info(f"{files} zipped into {str(zf_filepath_name['name'])}")

#---------------------------------------------------------------------------------------------
def zip_mult_fldr_files(folders, zip_filename):
    zip_file = zipfile.ZipFile(zip_filename, 'w',
                        compression=zipfile.ZIP_DEFLATED,
                        compresslevel=6)

    for directory in folders:
        for file_path in directory.iterdir():
            zip_file.write(file_path, arcname=file_path.name)
            
    #with zipfile.ZipFile(zip_filename, mode="r") as zip_file:
        #zip_file.printdir()
#---------------------------------------------------------------------------------------------

def email(email_subject:str, email_To:str, email_Cc:str, email_body, outfile_path:str, outfile_name, 
          attach:bool = False, attachment_files:list = [], email_image=False):
    """
    Function Actions:
    - Creates an unsent email.
    - Attaches files in attachment_files list if attach=True.
    - Saves as .eml file in outfile_path.
    - Prints and logs a message of email subject.
    """
    # Create the message
    msg = MIMEMultipart()
    msg['Subject'] = email_subject
    msg['To']      = email_To
    msg['Cc']      = email_Cc
    msg.add_header('X-Unsent', '1')
    if email_image != False:
        msg.attach(email_image)
    msg.attach(email_body)
    
    # attach all files in list to email
    if attach == True:
        for filename in attachment_files:
            attachment = open(filename, 'rb')
            part = MIMEBase("application", "octet-stream")
            part.set_payload(attachment.read())
            encoders.encode_base64(part)
            part.add_header('Content-Disposition','attachment', filename=os.path.basename(filename))
            msg.attach(part)

    outfile_name = os.path.join(outfile_path, outfile_name)
    outfile_name = PurePath(outfile_name)

    #save file to outfile_path
    with open(outfile_name, 'w') as outfile:
        gen = generator.Generator(outfile)
        gen.flatten(msg)
    
    logging.info(str(outfile_name) + ' created and saved.')

#---------------------------------------------------------------------------------------------

def create_word_doc(template, YYYY:str, content:dict, par, business_name:str):
    """
    Function Actions:
    - Create a word document from a template file.
    - Read in template.
    - Insert copyright year into the footer.
    - Insert any content into the document (loops through content dictionary and does a find and repalce).
    - Returns the document.
    """
    document = docx.Document(template)

    # insert copyright year into footer
    if 'copyright_year' in document.sections[0].footer.paragraphs[par].text:
        document.sections[0].footer.paragraphs[par].text = f"Copyright © {YYYY} {business_name}"

    # insert content into word doc
    for i in content:
        for paragraph in document.paragraphs:
            if paragraph.text.find(i)>=0:
                paragraph.text=paragraph.text.replace(i,content[i])
       
    return document

#---------------------------------------------------------------------------------------------

def save_word_doc(document, outfile_path:str, outfile_name:str):
    """
    Function Actions:
    - Takes a document of docx.Document format.
    - Saves it to the outfile_path location.
    - Prints a message of document name and location.
    """

    document.save(PurePath(outfile_path, outfile_name))

    print(f'{outfile_name} Word Document populated and saved to {outfile_path}.')

#-------------------------------------------------------------------------------------------------
def write_summary_txt_file(filename, message):  
    with open(filename, 'a') as file1:
        file1.write("\n")
        file1.write(message)
        file1.close()
        
#------------------------------------------------------------------------------------------------
def bulk_replace_str_xlsx(folderloc, filename, sheetname, findstr, replacestr):
    """
    Function Actions:
        - Bulk replaces string in an excel document.
    Requirements: 
        -import openpyxl
        -from openpyxl.utils.cell import get_column_letter   
    """
    #specify which excel workbook to use
    wb = openpyxl.load_workbook(PurePath(folderloc, filename))
    
    #return all sheets from workbook and specify sheet we want to use
    wb.sheetnames
    sheet = wb[sheetname]
    #set max number of columns and rows found in a specified sheet
    number_rows = sheet.max_row
    number_columns = sheet.max_column

    #dictionary of string to find and then replace
    replacement = {findstr: replacestr}

    #run through each column and each row and identify any cells that matches the target string. If so, replace it. (Does not work for substrings in a cell).
    for i in range(number_columns):
        for k in range(number_rows):
            cell = str(sheet[get_column_letter(i+1)+str(k+1)].value)
            for key in replacement.keys():
                if str(cell) == key:
                    newCell = replacement.get(key)
                    sheet[get_column_letter(i+1)+str(k+1)] = str(newCell)

    wb.save(PurePath(folderloc, filename))


 #-------------------------------------------------------------------------------------------------------------------------------------------------------------
# Hyperlink text in a word doc - used for the supporting products document
def add_hyperlink(paragraph, text, url):
    """
    Function Actions:
    - Turns text into a hyperlink, within a word document.
    """
    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element and a new w:rPr element
    new_run = docx.oxml.shared.OxmlElement('w:r')
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    # Create a new Run object and add the hyperlink into it
    r = paragraph.add_run()
    r._r.append (hyperlink)

    # A workaround for the lack of a hyperlink style (doesn't go purple after using the link)
    # Delete this if using a template that has the hyperlink style in it
    r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK

    return hyperlink

#-------------------------------------------------------------------------------------------------------------------------------------------------------------

def newline_character_found(df):
    #define character search filter 
    mask = np.column_stack([df[col].str.contains(r"\n", na=False) for col in df]) #dependent on all columns being str data type

    #filter for rows where any column contains 'Guard'
    rows_contain_newline = df.loc[mask.any(axis=1)] 
    rows_contain_newline

    #return the column where the substring is found (based off the row conditions)
    cols_contain_newline = rows_contain_newline.loc[:,df.gt('kr').any()]

    #make a slist of the affected columns that contain a hidden new line
    cols_contain_newline_list = cols_contain_newline.columns.values.tolist()
    cols_contain_newline_list

    #print logger warning message that details user needs to manually check and remove hidden new lines (\n) using a stored procedure...
    logging.warning('Hidden newlines characters have been identified in the SQL table. These must be removed to avoid affecting table creations in python. \n Please run xx stored procedure to amend this, affected table' +  'TABLE_link' + 'and specifically columns' + cols_contain_newline_list)


    f"""
    ---check where rows are affected
    SELECT TOP (1000) *
    FROM {ClusterManagement.Clusters}
    --WHERE Cluster_ID =  'UNPAIDCARER_COD'
    WHERE Cluster_Authorisation like '%' + CHAR(10) + '%'


    --make actual update to these rows, specifying all at a specific column
    UPDATE {ClusterManagement.Clusters}
    set Cluster_Category = REPLACE(Cluster_Category, CHAR(10), '')
    --WHERE Cluster_ID =  'DMNONTYPE1_COD'
    """

    #-------------------------------------------------------------------------------------------------------------------------------------------------------------